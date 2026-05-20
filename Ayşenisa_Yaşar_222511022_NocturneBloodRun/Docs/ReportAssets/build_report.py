from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ASSET_DIR = ROOT / "Docs" / "ReportAssets"
OUTPUT = ROOT / "Docs" / "Reports" / "Aysenisa_Yasar_222511022_Nocturne_Blood_Run_Ayrintili_Rapor.docx"

ACCENT = RGBColor(28, 49, 84)
ACCENT_SOFT = RGBColor(221, 232, 247)
SECONDARY = RGBColor(110, 76, 38)
TEXT = RGBColor(35, 40, 52)
MUTED = RGBColor(92, 100, 112)
GREEN = RGBColor(56, 111, 79)
RED = RGBColor(130, 52, 52)


def set_run_font(run, name="Calibri", size=11, bold=False, color=TEXT, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    no_proof = OxmlElement("w:noProof")
    run._element.rPr.append(no_proof)


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=10.5, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, after=2, line_spacing=1.05)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def add_paragraph(doc, text, size=11, bold=False, italic=False, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=before, after=after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_heading(doc, text, size=17, color=ACCENT, before=8, after=7):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, before=before, after=after)
    run = paragraph.add_run(text)
    set_run_font(run, name="Cambria", size=size, bold=True, color=color)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D6E2F2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return paragraph


def add_subheading(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=8, after=4)
    run = paragraph.add_run(text)
    set_run_font(run, name="Cambria", size=13.5, bold=True, color=SECONDARY)
    return paragraph


def add_bullets(doc, items, size=10.8):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, after=3, line_spacing=1.08)
        run = paragraph.add_run(item)
        set_run_font(run, size=size)


def add_numbered(doc, items, size=10.8):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(paragraph, after=3, line_spacing=1.08)
        run = paragraph.add_run(item)
        set_run_font(run, size=size)


def add_callout(doc, title, lines, fill="F4F7FB", title_color=ACCENT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    cell.width = Inches(6.4)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, after=5)
    title_run = paragraph.add_run(title + "\n")
    set_run_font(title_run, name="Cambria", size=12, bold=True, color=title_color)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, size=10.6)
        if index != len(lines) - 1:
            paragraph.add_run("\n")


def add_figure(doc, image_path, caption, width_inches=6.5):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=4, after=3)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption_paragraph, after=9)
    run = caption_paragraph.add_run(caption)
    set_run_font(run, size=10, italic=True, color=MUTED)


def build_cover(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=18, after=8)
    run = title.add_run("NOCTURNE BLOOD RUN\nAYRINTILI PROJE RAPORU")
    set_run_font(run, name="Cambria", size=22, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=14)
    run = subtitle.add_run("Unity Scripting - Visualization Basics Kapsaminda Teknik ve Gorsel Inceleme")
    set_run_font(run, size=12.5, bold=True, color=SECONDARY)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta.columns[0].width = Inches(2.1)
    meta.columns[1].width = Inches(4.1)
    meta_data = [
        ("Ogrenci", "Aysenisa Yasar"),
        ("Numara", "222511022"),
        ("Tarih", date(2026, 5, 20).strftime("%d.%m.%Y")),
        ("Proje", "Nocturne Blood Run"),
    ]
    for row_idx, (label, value) in enumerate(meta_data):
        left = meta.cell(row_idx, 0)
        right = meta.cell(row_idx, 1)
        set_cell_text(left, label, bold=True, size=11, color=RGBColor(255, 255, 255))
        set_cell_text(right, value, size=11)
        shade_cell(left, "1C3154")
        shade_cell(right, "F8F6F0")

    add_paragraph(
        doc,
        "Bu rapor, Unity ile gelistirilen Nocturne Blood Run projesinin teknik altyapisini, sahne kurgusunu, script organizasyonunu ve ders kapsaminda talep edilen Unity kavramlarini ayrintili bicimde incelemek amaciyla hazirlanmistir.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11.1,
        after=10,
    )

    add_callout(
        doc,
        "Raporun Amaci",
        [
            "Projede kullanilan her zorunlu Unity kavramini somut kod ornekleri ve dosya referanslariyla aciklamak.",
            "Oyunun sahne kurulumunu, oynanis akisini ve scriptler arasi veri akisini gorsellerle desteklemek.",
            "Uygulamanin neden bu yapida tasarlandigini ve hangi teknik secimlerin yapildigini degerlendirmek.",
        ],
        fill="EEF3FB",
    )

    add_figure(
        doc,
        ASSET_DIR / "scene_topdown_map.png",
        "Gorsel 1. Kasaba, orman, baslangic noktalar, altinlar ve kacis rotalarinin ustten sematik yerlesimi.",
        width_inches=6.7,
    )
    doc.add_page_break()


def ensure_additional_assets():
    gameplay_flow = ASSET_DIR / "gameplay_flow_diagram.png"
    if gameplay_flow.exists():
        return

    from PIL import Image, ImageDraw, ImageFont

    font_path = r"C:\Windows\Fonts\arial.ttf"
    font_bold = r"C:\Windows\Fonts\arialbd.ttf"
    title_font = ImageFont.truetype(font_bold, 34)
    head_font = ImageFont.truetype(font_bold, 22)
    body_font = ImageFont.truetype(font_path, 17)
    small_font = ImageFont.truetype(font_path, 15)

    width, height = 1600, 900
    image = Image.new("RGB", (width, height), "#f6f5ef")
    draw = ImageDraw.Draw(image)
    draw.text((60, 40), "Nocturne Blood Run - Oynanis Akis Diyagrami", font=title_font, fill="#1c3154")

    steps = [
        ("1. Baslangic", "Start paneli acilir ve oyuncu Baslat butonuna basar.", "#dce8fb"),
        ("2. Spawn", "Director oyunculari, canavari ve altinlari sahneye kurar.", "#e5f3df"),
        ("3. Kacis ve Toplama", "Oyuncular kasaba ve ormanda hareket ederek altin toplar.", "#fff1d9"),
        ("4. Silah Kilidi", "Takim skoru 50 oldugunda silah kullanimi acilir.", "#fce9d7"),
        ("5. Ates Etme", "Raycast ile canavara ates edilir, tracer ve hit efektleri cikar.", "#e8e1fb"),
        ("6. Sonuc", "Canavar oyuncuyu yerse oyun biter; canavar duserse zafer olur.", "#fde2e2"),
    ]

    positions = [(80, 160), (550, 160), (1020, 160), (80, 480), (550, 480), (1020, 480)]
    box_w = 360
    box_h = 180
    for (title, body, fill), (x, y) in zip(steps, positions):
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=24, fill=fill, outline="#6f84a4", width=4)
        draw.text((x + 24, y + 22), title, font=head_font, fill="#1f3557")
        draw.multiline_text((x + 24, y + 70), body, font=body_font, fill="#273142", spacing=6)

    arrows = [
        ((440, 250), (550, 250)),
        ((910, 250), (1020, 250)),
        ((1200, 340), (1200, 480)),
        ((1020, 570), (910, 570)),
        ((440, 570), (550, 570)),
    ]
    for start, end in arrows:
        draw.line((start, end), fill="#58739c", width=8)
        if end[0] > start[0]:
            tip = [(end[0], end[1]), (end[0] - 20, end[1] - 12), (end[0] - 20, end[1] + 12)]
        elif end[0] < start[0]:
            tip = [(end[0], end[1]), (end[0] + 20, end[1] - 12), (end[0] + 20, end[1] + 12)]
        else:
            tip = [(end[0], end[1]), (end[0] - 12, end[1] - 20), (end[0] + 12, end[1] - 20)]
        draw.polygon(tip, fill="#58739c")

    note = (
        "Akis mantigi: butonla baslayan oyun once spawn ve kontrol asamasina gecer, "
        "ardindan altin toplama ve silah kilidi ile ilerler. Canavarin yenilmesi ya da "
        "oyuncularin yenilmesi son durumu belirler."
    )
    draw.multiline_text((90, 770), note, font=small_font, fill="#48586c", spacing=5)
    image.save(gameplay_flow)


def add_executive_summary(doc):
    add_heading(doc, "1. Yonetici Ozeti")
    add_paragraph(
        doc,
        "Nocturne Blood Run, Unity uzerinde gelistirilmis iki oynanabilir karakterli bir kacis ve hayatta kalma prototipidir. Oyunun ana fikri, karanlik bir kasaba ve onu cevreleyen sonbahar ormaninda ilerleyerek altin toplamak, yeterli skor sonrasinda silah kilidini acmak ve oyunculari avlayan canavari etkisiz hale getirmektir.",
        after=5,
    )
    add_paragraph(
        doc,
        "Proje yalnizca oynanis mantigini degil, sahne kurulumunu da otomatiklestiren bir edit or araci ile olusturulmustur. Bu sayede ortam prefablarindan animasyon denetleyicilerine, NavMesh kurulumundan HUD ve buton yapisina kadar pek cok ogenin tekrar uretilebilir olmasi saglanmistir.",
        after=6,
    )
    add_bullets(
        doc,
        [
            "Remy ve Peasant Girl adli iki karakter, farkli kontrol semalari ile ayni sahnede ayni anda yonetilebilir.",
            "Canavar, NavMeshAgent ve gorus tabanli RayCast kontrolu ile oyunculari kovalar.",
            "Altin toplama, skor sistemi, silah kilidi, ates etme, oyun sonu ve yeniden baslatma akislari tek bir Director scripti tarafindan koordine edilir.",
            "SceneBuilder tarafinda uretilen prefablar ve sahne elemanlari, raporun merkezindeki Unity kavramlarinin buyuk cogunlugunu dogrudan icermektedir.",
        ],
    )
    add_callout(
        doc,
        "Bu Raporda Neler Var?",
        [
            "Projenin konusu ve sahne dili",
            "Script mimarisi ve veri akis diyagrami",
            "10 zorunlu Unity kavraminin tamamina ait ayrintili analiz",
            "Dosya ve satir referanslari ile teknik degerlendirme",
        ],
        fill="F6F8FC",
    )


def add_project_context(doc):
    add_heading(doc, "2. Projenin Konusu, Hedefi ve Oynanis Kurgusu")
    add_paragraph(
        doc,
        "Oyunda iki insan karakter, canavarin saldirisindan kurtulmak icin kasabanin ana yolunu, plaza alanini ve orman yollarini kullanarak hareket eder. Baslangicta oyuncular sadece kacis ve altin toplama davranisina sahiptir. Takim skoru gerekli esige ulastiginda ise ates etme yetenegi acilir ve oyun savunmadan saldiriya dogru evrilir.",
    )
    add_paragraph(
        doc,
        "Bu kurgu, ders kapsaminda beklenen birden fazla Unity kavramini tek bir hikaye icerisinde toplamaya imkan vermistir. Altinlar Trigger ile toplanmakta, oyuncu hareketi NavMesh uzerinde ilerlemekte, canavar RayCast ile hedefini gormekte ve partikuller hem ortam hem de etkilesim geri bildirimi icin kullanilmaktadir.",
    )
    add_numbered(
        doc,
        [
            "Oyun butonla baslar ve Director senaryoyu aktif eder.",
            "Oyuncular WASD ve yon tuslari ile hareket ederek altin toplar.",
            "Skor 50 oldugunda silah kullanimi acilir.",
            "Oyuncular RayCast tabanli ates ile canavari hasarlayabilir.",
            "Canavar oyuncuya yetisirse oyun biter; canavar dusurulurse zafer ekrani olusur.",
        ],
    )
    add_figure(
        doc,
        ASSET_DIR / "gameplay_flow_diagram.png",
        "Gorsel 2. Oyunun baslangictan sonuca kadar ilerleyen temel oynanis akisi.",
        width_inches=6.7,
    )


def add_visual_design(doc):
    add_heading(doc, "3. Gorsel Tasarim ve Sahne Cozumlemesi")
    add_paragraph(
        doc,
        "Sahne tasarimi, karanlik bir kasaba ile onu cevreleyen sonbahar ormaninin karsitligi uzerine kurulmustur. Merkezde yol ve plaza gibi daha okunakli geometri kullanilirken, cevredeki orman daha yogun ağaç tekrarları ile kapatilmistir. Bu yapinin amaci, hem kacis rotalarini netlestirmek hem de tehdit hissini arttirmaktir.",
    )
    add_bullets(
        doc,
        [
            "Kasaba merkezi: ana yol, plaza, evler, lambalar ve kapidan olusan kontrollu oynanis omurgasi.",
            "Orman katmani: yan kacis imkani sunan ancak gorus ve takip baskisini arttiran bir cevre.",
            "Işık duzeni: karanlik tonlu arka plan uzerine mavi-gri moonlight ve sicak lamba isiklari.",
            "Partikul atmosferi: dusuk yogunluklu yaprak suruklenmesi ve zemin sisi ile sonbahar havasi.",
        ],
    )
    add_figure(
        doc,
        ASSET_DIR / "scene_topdown_map.png",
        "Gorsel 3. Oyun alani; kasaba merkezi, orman bantlari, altin noktalar ve canavar baslangici ile birlikte sematik olarak gorulmektedir.",
        width_inches=6.6,
    )
    add_callout(
        doc,
        "Sahne Tasariminda One Cikan Teknik Secimler",
        [
            "SceneBuilder yolu, curb elemanlarini, evleri ve orman nesnelerini editorde otomatik uretir.",
            "Altin noktalarinin elde tek tek yerlestirilmesi yerine Vector3[] dizisi ile toplu olusturma tercih edilmistir.",
            "Atmosfer efektleri yolun tam ortasini kapatmayacak sekilde kenarlara ve kritik noktalara dagitilmistir.",
        ],
        fill="F7F3EA",
        title_color=SECONDARY,
    )


def add_architecture_section(doc):
    add_heading(doc, "4. Script Mimarisi ve Veri Akisi")
    add_paragraph(
        doc,
        "Projede sorumluluklar farkli scriptlere bolunmustur. Bu dagilim, kodun okunabilirligini ve yeni ozellik ekleme kolayligini arttirir. Ozetle SceneBuilder editorde sahneyi uretir; NocturneScenarioDirector oyun aninda senaryoyu baslatir; SurvivorAgent ve MonsterAgent bireysel davranislari yonetir; CameraRigFollow ve ShotTracer ise yardimci goruntu katmanini destekler.",
    )
    add_figure(
        doc,
        ASSET_DIR / "architecture_diagram.png",
        "Gorsel 4. Script bazli sorumluluk dagilimi ve oyun icindeki veri akisinin ozet diyagrami.",
        width_inches=6.7,
    )

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(2.2)
    table.columns[2].width = Inches(2.4)
    headers = ["Dosya", "Ana Sorumluluk", "Dikkat Ceken API veya Yapi"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=10.5, color=RGBColor(255, 255, 255))
        shade_cell(table.rows[0].cells[index], "1C3154")

    rows = [
        ("NocturneScenarioDirector.cs", "Spawn, HUD, skor ve oyun durumu yonetimi", "Instantiate, arrays, buton akisi"),
        ("SurvivorAgent.cs", "Oyuncu hareketi, ates etme, animator surumu", "FixedUpdate, RaycastAll, Quaternion"),
        ("MonsterAgent.cs", "Takip, gorus kontrolu ve saldiri", "NavMeshAgent, Raycast, Animator Trigger"),
        ("CameraRigFollow.cs", "Kameranin hedeflere gore hareketi", "LateUpdate, deltaTime, Transform[]"),
        ("GoldCollectible.cs", "Altin animasyonu ve toplama", "OnTriggerEnter, Instantiate"),
        ("EscapeZone.cs", "Kacis alanina giris kontrolu", "OnTriggerEnter"),
        ("ShotTracer.cs", "Mermi izinin fade edilmesi", "LineRenderer, Time.deltaTime"),
        ("NocturneVillageSceneBuilder.cs", "Prefab ve sahne otomasyonu", "PrefabUtility, AnimatorController, BuildNavMesh"),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0], bold=True, size=9.9)
        set_cell_text(cells[1], row[1], size=9.9)
        set_cell_text(cells[2], row[2], size=9.9)
        for cell in cells:
            shade_cell(cell, "FBFAF6")


def add_compliance_table(doc):
    add_heading(doc, "5. Ders Kriterleri ve Karsilanma Durumu")
    add_paragraph(
        doc,
        "Asagidaki tablo, ders gereksinimlerinde istenen tum Unity basliklarinin projede karsilandigini ve bunlarin hangi scriptlerde yer aldigini ozetlemektedir.",
        after=6,
    )
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.6), Inches(1.0), Inches(2.2), Inches(2.0)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    headers = ["Baslik", "Durum", "Ana Dosyalar", "Not"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=10.2, color=RGBColor(255, 255, 255))
        shade_cell(table.rows[0].cells[idx], "1C3154")

    rows = [
        ("Prefab", "Var", "SceneBuilder, Director", "Karakter, ortam, efekt ve gameplay nesneleri prefab olarak kurulmustur."),
        ("Vector3 / Quaternion", "Var", "SurvivorAgent, CameraRigFollow, SceneBuilder", "Hareket, bakis, mermi cikisi ve sahne yerlesimi bu yapiyla hesaplanir."),
        ("Partikul efektleri", "Var", "SceneBuilder, GoldCollectible, MonsterAgent", "Spawn, vurulma, altin ve atmosfer efektleri vardir."),
        ("Instantiate", "Var", "Director, SurvivorAgent, EscapeZone", "Runtime nesne uretimi aktif sekilde kullanilir."),
        ("RayCast", "Var", "SurvivorAgent, MonsterAgent", "Hem ates hem de canavar gorusu icin kullanilir."),
        ("FixedUpdate / LateUpdate", "Var", "SurvivorAgent, CameraRigFollow", "Fizik ve kamera mantigi farkli update dongulerine ayrilmistir."),
        ("Trigger / Collision", "Var", "GoldCollectible, EscapeZone", "Trigger tabanli etkileşim tercih edilmistir."),
        ("Arrays", "Var", "Director, CameraRigFollow, SceneBuilder", "Spawn noktasi ve hedef listeleri dizilerle tutulur."),
        ("NavMeshAgent", "Var", "SurvivorAgent, MonsterAgent, SceneBuilder", "Takip ve hareketin temel yol bulma katmanidir."),
        ("Animator", "Var", "SceneBuilder, SurvivorAgent, MonsterAgent", "Idle, run, walk ve attack gecisleri kontrol edilir."),
    ]
    for title, status, files, note in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], title, bold=True, size=9.8)
        set_cell_text(cells[1], status, bold=True, size=9.8, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[2], files, size=9.8)
        set_cell_text(cells[3], note, size=9.6)
        for cell in cells:
            shade_cell(cell, "F9FAFC")


def add_concept_detail(doc, index, title, definition, usage, references, gains, caution=None):
    add_subheading(doc, f"6.{index} {title}")
    add_paragraph(doc, "Kavramsal aciklama:", size=10.8, bold=True, color=ACCENT, after=3)
    add_paragraph(doc, definition, size=10.7, after=4)
    add_paragraph(doc, "Projede nasil kullanildi:", size=10.8, bold=True, color=ACCENT, after=3)
    add_paragraph(doc, usage, size=10.7, after=4)
    add_paragraph(doc, "Kod referanslari:", size=10.8, bold=True, color=ACCENT, after=3)
    add_bullets(doc, references, size=10.3)
    add_paragraph(doc, "Projeye katkisi:", size=10.8, bold=True, color=ACCENT, after=3)
    add_bullets(doc, gains, size=10.3)
    if caution:
        add_callout(doc, "Teknik Not", [caution], fill="F8F3EC", title_color=SECONDARY)


def add_concepts_section(doc):
    add_heading(doc, "6. Zorunlu Unity Kavramlarinin Ayrintili Analizi")
    add_paragraph(
        doc,
        "Bu bolumde, ders listesinde istenen on temel Unity basligi tek tek incelenmektedir. Her kavram icin once ne oldugu aciklanmakta, daha sonra Nocturne Blood Run projesinde hangi kod bloklarinda kullanildigi ve bu secimin oynanisa ne kattigi anlatilmaktadir.",
    )

    details = [
        (
            "Prefab Kullanimi",
            "Prefab, bir oyun nesnesinin ayarlariyla birlikte tekrar kullanilabilir sablon olarak saklanmasidir. Unity projelerinde ayni tip nesneyi tekrar tekrar kullanirken tutarlilik ve hiz saglar.",
            "Projede prefab mantigi iki katmanda kullanilmistir: birincisi SceneBuilder tarafinda karakter, ev, agac, lamba, altin ve efekt prefablarinin uretilmesi; ikincisi Director tarafinda bu prefablarin runtime sirasinda Instantiate edilerek sahneye yerlestirilmesidir.",
            [
                "NocturneVillageSceneBuilder.cs 358-381: efekt ve gameplay prefablarinin olusturulmasi.",
                "NocturneVillageSceneBuilder.cs 426-528: survivor ve monster prefablarinin Animator, NavMeshAgent ve gameplay scriptleriyle birlestirilmesi.",
                "NocturneScenarioDirector.cs 199-238: prefablarin oyun aninda spawn edilmesi.",
            ],
            [
                "Ayni nesne yapisinin birden fazla noktada tutarli bicimde kullanilmasini saglar.",
                "Sahne kurulumunu hizlandirir ve script tabanli otomasyonun temelini olusturur.",
                "Karakter veya efekt degisikligi gerektiğinde tek bir kaynaktan guncelleme imkani verir.",
            ],
            "Bu projede prefab kullanimi sadece nesne saklamak icin degil, sahneyi otomatik ureten edit or akisini da besleyen ana mekanizma olarak secilmistir.",
        ),
        (
            "Vector3 ve Quaternion Kullanimi",
            "Vector3, Unity sahnesindeki konum, yon ve olcek hesaplari icin kullanilan temel veri yapisidir. Quaternion ise acisal donusleri gimbal lock sorunu olmadan yonetmek icin kullanilir.",
            "Oyuncu hareketi, kameranin hedeflere bakmasi, mermi cikis noktasi, altin ve spawn pozisyonlari ile sahnedeki yol-ev-orman yerlesimleri bu iki yapiyla kurulmustur.",
            [
                "SurvivorAgent.cs 100-126: hareket yonunun Vector3 ile hesaplanmasi ve Quaternion.Slerp ile donus.",
                "SurvivorAgent.cs 303-309: mermi cikisi icin origin ve hitPoint hesaplari.",
                "CameraRigFollow.cs 58-63: desiredPosition ve desiredRotation hesaplari.",
                "NocturneVillageSceneBuilder.cs 602-643, 736-911: sahne nesnelerinin koordinat ve rotasyonlari.",
            ],
            [
                "Oyuncu hareketinin yonsel olarak dogru ve akici olmasini saglar.",
                "Kamera takibinin merkez noktasina gore stabil davranmasina yardim eder.",
                "Sahne icindeki tum mekansal organizasyonu script ile tekrar uretilebilir hale getirir.",
            ],
            None,
        ),
        (
            "Partikul Efektleri",
            "Partikul sistemleri, oyundaki olaylari gorsel olarak vurgulamak ve atmosfere hareket katmak icin kullanilir. Patlama, duman, sis, yaprak ve isik serpintisi gibi pek cok kucuk efekti verimli sekilde uretir.",
            "Bu projede spawn sirasinda mist, vurulma sirasinda impact ve blood burst, ates sirasinda muzzle flash, altin toplanirken gold burst, ortamda ise autumn drift ve ground fog kullanilmistir.",
            [
                "NocturneVillageSceneBuilder.cs 358-372: tum efekt prefablarinin uretildigi bolum.",
                "NocturneVillageSceneBuilder.cs 1614-1703: sonbahar yapragi ve zemin sisi gibi atmosfer efektlerinin ayarlanmasi.",
                "GoldCollectible.cs 52-54: altin pickup efekti.",
                "MonsterAgent.cs 220 ve 255: saldiri ve olum anlarindaki efektler.",
            ],
            [
                "Oyuncuya olaylarin sonucunu aninda anlatir.",
                "Karanlik kasaba temasini gorsel olarak destekler.",
                "Spawn, ates ve hasar gibi olaylarin geri bildirim kalitesini artirir.",
            ],
            "Partikuller yolun ortasinda gorusu kapatmayacak bicimde konumlandirilmis; bu secim oyun okunurlugunu korumak icin yapilmistir.",
        ),
        (
            "Instantiate Kullanimi",
            "Instantiate, bir prefab veya nesnenin calisma aninda yeni bir kopyasini uretmek icin kullanilir. Dinamik oyunlarda sabit sahne kurulumuna bagli kalmadan yeni nesne yaratmaya imkan verir.",
            "Projede oyuncular, canavar, altinlar, mermi izi ve ani efektlerin tamami runtime uretim mantigi ile olusturulur. Bu sayede oyun tekrar basladiginda sahne tekrar kurulabilir.",
            [
                "NocturneScenarioDirector.cs 199-238: survivor, canavar ve altinlarin olusturulmasi.",
                "SurvivorAgent.cs 309-339: muzzle flash, impact burst ve shot tracer nesnelerinin olusturulmasi.",
                "EscapeZone.cs 35 ve GoldCollectible.cs 54: etkileşim anlarinda efekt uretimi.",
            ],
            [
                "Oynanisin durumuna gore nesne uretimine izin verir.",
                "Sahneyi tek seferlik degil, tekrar oynanabilir hale getirir.",
                "Efektlerin yalnizca gerektiginde ortaya cikip daha hafif bir sahne akisi saglamasina yardim eder.",
            ],
            None,
        ),
        (
            "RayCast Kullanimi",
            "RayCast, sahnede bir dogrultuda isik benzeri cizgi gonderip ilk carpisma noktasini tespit etmeye yarar. Atis mekanikleri, hedef algilama ve gorus kontrolu gibi durumlarda oldukca yaygindir.",
            "Projede oyuncularin ates sistemi RaycastAll ile, canavarin oyuncuyu gorme mantigi ise Raycast ile kurulmustur. Boylece hem ofansif mekanik hem de AI algilama ayni Unity mantigiyle desteklenmistir.",
            [
                "SurvivorAgent.cs 312-326: RaycastAll ile merminin carptigi ilk gecerli hedefin bulunmasi.",
                "SurvivorAgent.cs 323-326: vurulan collider icinden MonsterAgent referansinin alinmasi.",
                "MonsterAgent.cs 178-182: canavarin dogrudan oyuncuyu gorup gormedigini test eden Physics.Raycast.",
            ],
            [
                "Ates mekanigine fizik temelli bir mantik kazandirir.",
                "Canavarin her zaman degil, yalnizca gorus alabildiginde agresif davranmasini saglar.",
                "Sahne engellerinin oynanisa etkide bulunmasina imkan verir.",
            ],
            "RayCast kullanimi, oyundaki tehdit algisini guclendirir cunku canavar sadece mesafe degil, gorus dogrulamasi da yapmaktadir.",
        ),
        (
            "FixedUpdate, LateUpdate ve DeltaTime",
            "Unity'de farkli update donguleri farkli sorumluluklar icin tercih edilir. FixedUpdate daha kararlı fizik mantigi sunar; LateUpdate ise diger nesneler hareket ettikten sonra kamera gibi bagimli sistemleri guncellemek icin uygundur. DeltaTime ve fixedDeltaTime hareketleri kare hizindan bagimsiz yapar.",
            "Projede oyuncu hareketi FixedUpdate icine alinarak agent.Move hesaplari sabit zaman adimiyla yapilmistir. Kamera ise LateUpdate icinde karakterlerin son pozisyonlarini izler. Altin animasyonu ve tracer fade davranisi Time.deltaTime ile yumusatilmiştir.",
            [
                "SurvivorAgent.cs 93-126: FixedUpdate ile fizik uyumlu hareket ve donus.",
                "CameraRigFollow.cs 19-63: LateUpdate ile kamera merkezleme ve bakis hesaplari.",
                "GoldCollectible.cs 26-27: delta time kullanarak donme ve bobbing animasyonu.",
                "ShotTracer.cs 18-33: elapsed += Time.deltaTime ile iz efekti sönmesi.",
            ],
            [
                "Farkli sistemlerin dogru zamanda guncellenmesini saglar.",
                "Kare hizi degisse bile hareketin tutarliligini korur.",
                "Kameranin oyuncudan geri kalmadan pürüzsüz izleme yapmasina yardim eder.",
            ],
            None,
        ),
        (
            "Trigger Kullanimi",
            "Trigger collider, fiziksel carpismadan ziyade bir alanin icine girildigini tespit etmek icin kullanilir. Oyunlarda toplama, bolgeye giris ve gorev tetikleme gibi senaryolar icin idealdir.",
            "Bu projede altin toplama ve kacis alani mantiklari OnTriggerEnter uzerinden kurulmustur. Collision tercih edilmemistir cunku burada fiziksel itme degil, olay tetikleme davranisi istenmektedir.",
            [
                "GoldCollectible.cs 17-21: SphereCollider icin isTrigger ayari.",
                "GoldCollectible.cs 37-58: altina dokunan survivor icin toplama akisi.",
                "EscapeZone.cs 17-20: bolgenin trigger olarak ayarlanmasi.",
                "EscapeZone.cs 23-35: survivor kacis alanina girince ReachSafety akisinin cagrilmasi.",
            ],
            [
                "Toplama ve bolgesel etkileşimleri fizik yiginina gereksiz yuk bindirmeden cozer.",
                "Oyuncu hareketini bozmadan olay bazli tetikleme yapar.",
                "Skor, gorev ve oyun sonu mantiklariyla temiz sekilde baglanir.",
            ],
            "Bu projede Collision yerine Trigger secilmesi bilincli bir tasarim kararidir; cunku altinlar oyuncuyu itmemeli, sadece algilanmalidir.",
        ),
        (
            "Arrays Kullanimi",
            "Diziler, ayni tipte birden fazla veriyi sabit sirali yapida saklamaya yarar. Oyun gelistirmede spawn noktalarini, hedef listelerini ve toplu islenecek nesneleri yonetmek icin verimli bir cozumdur.",
            "Projede oyuncu prefab dizileri, spawn noktasi dizileri, altin noktasi dizileri ve kamera hedef dizileri sistematik bir kurulum mantigi olusturur. SceneBuilder tarafinda da Vector3[] dizileriyle ortam yerlestirmesi yapilmaktadir.",
            [
                "NocturneScenarioDirector.cs 10-27: survivorPrefabs, survivorSpawnPoints, goldSpawnPoints ve spawnedSurvivors tanimlari.",
                "NocturneScenarioDirector.cs 193-243: diziler uzerinden dongu ile senaryo kurulumu.",
                "CameraRigFollow.cs 12-16: takip hedeflerinin Transform[] olarak atanmasi.",
                "NocturneVillageSceneBuilder.cs 630-643 ve 905-927: altin, drift ve fog pozisyon dizileri.",
            ],
            [
                "Tek tek nesne atamak yerine toplu yonetim saglar.",
                "Sahne buyudukce bile kodun duzenli kalmasina yardim eder.",
                "Yeni spawn noktasi eklemeyi veya mevcutlari degistirmeyi kolaylastirir.",
            ],
            None,
        ),
        (
            "Navigation Mesh ve NavMeshAgent Kullanimi",
            "NavMesh, sahnede yurunebilir alanlari temsil eden bir navigasyon katmanidir. NavMeshAgent ise bu alan uzerinde hareket eden ve engelleri dolasan yapay zeka veya karakter bilesenidir.",
            "Projede hem survivor hem de canavar prefablarina NavMeshAgent eklenmistir. Survivor tarafinda hareket dogrudan agent.Move ile surulur; canavar tarafinda ise SetDestination kullanilarak hedef takip edilir.",
            [
                "NocturneVillageSceneBuilder.cs 450-456: survivor prefab agent ayarlari.",
                "NocturneVillageSceneBuilder.cs 499-505: monster prefab agent ayarlari.",
                "NocturneVillageSceneBuilder.cs 716: BuildNavMesh ile sahnenin navigasyon verisinin olusturulmasi.",
                "SurvivorAgent.cs 122: agent.Move ile kontrollu hareket.",
                "MonsterAgent.cs 97-103: SetDestination ile hedef kovalamasi.",
            ],
            [
                "Canavarin kasaba ve orman icinde gercekci sekilde yol bulmasini saglar.",
                "Oyuncu hareketini grid bagimli olmaktan cikarir ve sahne topolojisine uyarlar.",
                "Sahne buyuk oldugu icin manuel rota mantigi yazma ihtiyacini azaltir.",
            ],
            "NavMesh kullanimi, orman ve kasaba gibi farkli geometri alanlarinda davranis kararliligini belirgin bicimde arttirmistir.",
        ),
        (
            "Animation ve Animator Kullanimi",
            "Animator sistemi, bir karakterin farkli animasyon klipleri arasinda kosullu gecisler yapmasini saglar. Hangi animasyonun ne zaman oynatilacagi parametrelerle kontrol edilir.",
            "Projede survivor karakterleri icin Idle ve Run gecisleri, canavar icin ise Idle, Walk ve Attack durumlari tanimlanmistir. Hareket hizina gore Speed parametresi guncellenir; saldiri aninda Attack trigger'i ateslenir.",
            [
                "NocturneVillageSceneBuilder.cs 270-335: Animator Controller ve state machine olusturma.",
                "NocturneVillageSceneBuilder.cs 443-445 ve 492-494: runtime animator controller baglantilari.",
                "SurvivorAgent.cs 353-360: Speed parametresinin hereket buyuklugune gore guncellenmesi.",
                "MonsterAgent.cs 206-208 ve 269: Attack trigger ve Speed set islemleri.",
            ],
            [
                "Karakterlerin dururken, yururken ve kosarken dogal gorunmesini saglar.",
                "Canavarin saldıri aninin gorsel olarak anlasilir olmasina yardim eder.",
                "Oyundaki hareket verisini gorsel ifade ile birlestirir.",
            ],
            None,
        ),
    ]

    for index, detail in enumerate(details, start=1):
        add_concept_detail(doc, index, *detail)
        if index in (3, 6, 8):
            doc.add_page_break()


def add_runtime_flow_section(doc):
    add_heading(doc, "7. Oyun Akisinin Teknik Yorumlanmasi")
    add_paragraph(
        doc,
        "Nocturne Blood Run sadece kavram listelerini gosteren bir demo degil, birbirine bagli bir oynanis sistemi olarak tasarlanmistir. Aşağıdaki akis oyunun calisma mantigini teknik acidan yorumlamaktadir.",
    )
    add_numbered(
        doc,
        [
            "Oyun acildiginda Director sahneyi hazirlar fakat karakterleri hareketsiz tutar; start paneli oyuncudan buton girdisi bekler.",
            "BeginGame cagrildiginda survivor scriptleri aktif hale gelir ve canavar davranisi acilir.",
            "Oyuncular altin topladikca RegisterGoldPickup skor sistemini gunceller ve gerekli esikte silah kilidini acar.",
            "Silah acildiktan sonra SurvivorAgent Raycast tabanli ates sistemini kullanarak canavara hasar verebilir.",
            "MonsterAgent hedefi buldugunda SetDestination ve Attack coroutine akisi ile saldiri davranisina gecer.",
            "Survivor yenilirse Director oyun sonu durumunu yazar; canavar yenilirse zafer akisi tetiklenir.",
        ],
    )
    add_callout(
        doc,
        "Akis Tasarimindaki Onemli Nokta",
        [
            "Sistemin merkezi Director olsa da hareket, gorus, saldiri, kamera ve toplama gibi alt davranislar ilgili scriptlere dagitilmistir.",
            "Bu dagilim yeni mekanik eklerken mevcut yapinin bozulmamasini kolaylastirir.",
        ],
        fill="EEF6EE",
        title_color=GREEN,
    )


def add_code_reference_section(doc):
    add_heading(doc, "8. Dosya ve Kod Referans Ozeti")
    add_paragraph(
        doc,
        "Asagidaki tablo, raporda en cok referans verilen dosyalarin hangi gorevi ustlendiklerini ve inceleme sirasinda neden on plana ciktigini ozetler.",
        after=6,
    )
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(2.0), Inches(1.3), Inches(1.7), Inches(2.0)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    headers = ["Dosya", "Katman", "Temel Sorumluluk", "Rapor Icindeki Onemi"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=10.2, color=RGBColor(255, 255, 255))
        shade_cell(table.rows[0].cells[idx], "1C3154")

    rows = [
        ("Assets/Editor/NocturneVillageSceneBuilder.cs", "Editor", "Prefab, sahne, NavMesh ve HUD uretimi", "Zorunlu kavramlarin buyuk kismi editorde kurulan altyapiyla iliskilidir."),
        ("Assets/Scripts/Nocturne/NocturneScenarioDirector.cs", "Runtime", "Spawn, skor, butonlar, oyun sonu", "Tüm oynanis akisini koordine eder."),
        ("Assets/Scripts/Nocturne/SurvivorAgent.cs", "Runtime", "Oyuncu kontrolu ve ates", "Vector3, Quaternion, FixedUpdate, RayCast ve Animator burada bir araya gelir."),
        ("Assets/Scripts/Nocturne/MonsterAgent.cs", "Runtime", "Takip ve saldiri AI", "NavMeshAgent, RayCast ve Animator davranislarini birlestirir."),
        ("Assets/Scripts/Nocturne/CameraRigFollow.cs", "Runtime", "Kamera takibi", "LateUpdate ve deltaTime kullanimini gosteren net bir ornektir."),
        ("Assets/Scripts/Nocturne/GoldCollectible.cs", "Runtime", "Altin toplama", "Trigger, Instantiate ve animasyon benzeri gorsel davranislari icerir."),
        ("Assets/Scripts/Nocturne/EscapeZone.cs", "Runtime", "Kacis alaninin tetiklenmesi", "Trigger mantigini ikinci bir ornekle destekler."),
        ("Assets/Scripts/Nocturne/ShotTracer.cs", "Runtime", "Ates izinin fade edilmesi", "deltaTime ile gorsel efekt omrunu yonetir."),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0], bold=True, size=9.2)
        set_cell_text(cells[1], row[1], size=9.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[2], row[2], size=9.2)
        set_cell_text(cells[3], row[3], size=9.2)
        for cell in cells:
            shade_cell(cell, "FCFCFA")


def add_evaluation_section(doc):
    add_heading(doc, "9. Teknik Degerlendirme ve Sonuc")
    add_paragraph(
        doc,
        "Nocturne Blood Run projesi, ders kapsaminda beklenen Unity konu basliklarini tek tek gostermek yerine bunlari bir oyun dongusu icinde birlestirdigi icin daha guclu bir uygulama ornegidir. Bir yanda editorde sahne ureten bir otomasyon katmani, diger yanda runtime davranislari yoneten scriptler bulunmaktadir. Bu ayirim, projenin hem teknik olarak okunabilir hem de genisletilebilir olmasini saglar.",
    )
    add_bullets(
        doc,
        [
            "Projenin en guclu tarafi, prefab-temelli sahne kurulumu ile runtime mekaniklerini ayni yapida birlestirmesidir.",
            "NavMeshAgent, RayCast ve Animator secimleri, canavar ile oyuncular arasindaki etkileşimi net ve anlaşilir hale getirir.",
            "Trigger kullanim karari, altin toplama ve kacis alani gibi etkileşimlerde gereksiz fizik karmasasini engeller.",
            "Partikul efektleri ve atmosfer nesneleri, projenin gorsel dilini sadece teknik degil sunumsal olarak da guclendirir.",
        ],
    )
    add_callout(
        doc,
        "Genel Sonuc",
        [
            "Bu proje, Unity'nin temel script ve gorsellestirme kavramlarini ders seviyesinde karsilamanin otesine gecerek bir mini oyun sistemi halinde sunmaktadir.",
            "Dolayisiyla raporlanan her baslik teorik ornek olmaktan cikmis, sahnede calisan bir mekanige donusmustur.",
        ],
        fill="EDF4FE",
    )


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        run = paragraph.add_run("Aysenisa Yasar - 222511022 - Nocturne Blood Run Ayrintili Proje Raporu")
        set_run_font(run, size=9, color=MUTED)


def ensure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    if "CaptionCustom" not in doc.styles:
        style = doc.styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10)


def main():
    ensure_additional_assets()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    ensure_styles(doc)
    build_cover(doc)
    add_executive_summary(doc)
    add_project_context(doc)
    add_visual_design(doc)
    doc.add_page_break()
    add_architecture_section(doc)
    add_compliance_table(doc)
    doc.add_page_break()
    add_concepts_section(doc)
    add_runtime_flow_section(doc)
    add_code_reference_section(doc)
    add_evaluation_section(doc)
    add_footer(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
